import os
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import mm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
import markdown


class LessonPlanExporter:
    def __init__(self):
        self.setup_fonts()
    
    def setup_fonts(self):
        """Настройка шрифтов Times New Roman из папки проекта"""
        try:
            font_base_path = os.path.join(settings.BASE_DIR, 'app', 'static', 'fonts')
            
            times_roman_path = os.path.join(font_base_path, 'times.ttf')
            times_bold_path = os.path.join(font_base_path, 'timesbd.ttf')
            times_italic_path = os.path.join(font_base_path, 'timesi.ttf')
            times_bold_italic_path = os.path.join(font_base_path, 'timesbi.ttf')
            
            if os.path.exists(times_roman_path):
                pdfmetrics.registerFont(TTFont('Times-Roman', times_roman_path))
                print("Times-Roman font registered successfully")
            else:
                print(f"Times-Roman font not found at: {times_roman_path}")
            
            if os.path.exists(times_bold_path):
                pdfmetrics.registerFont(TTFont('Times-Bold', times_bold_path))
                print("Times-Bold font registered successfully")
            else:
                print(f"Times-Bold font not found at: {times_bold_path}")
                
            if os.path.exists(times_italic_path):
                pdfmetrics.registerFont(TTFont('Times-Italic', times_italic_path))
            
            if os.path.exists(times_bold_italic_path):
                pdfmetrics.registerFont(TTFont('Times-BoldItalic', times_bold_italic_path))
                
        except Exception as e:
            print(f"Font setup error: {e}")
    
    def markdown_to_plain_text(self, text):
        """Конвертирует Markdown в простой текст"""
        if not text:
            return ""
        
        text = text.replace('**', '').replace('*', '').replace('##', '').replace('###', '')
        text = text.replace('- ', '• ')
        
        html = markdown.markdown(text)
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    def create_pdf(self, lesson_plan):
        """Создает PDF файл плана урока"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              topMargin=20*mm, bottomMargin=20*mm,
                              leftMargin=20*mm, rightMargin=20*mm)
        
        styles = self.get_pdf_styles()
        story = []
        
        title_style = styles['Title']
        title = Paragraph(f"План урока: {lesson_plan.topic}", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        info_style = styles['Normal']
        info_text = f"""
        <b>Класс:</b> {lesson_plan.study_class.name}<br/>
        <b>Предмет:</b> {lesson_plan.study_class.subject}<br/>
        <b>Количество часов:</b> {lesson_plan.hours}<br/>
        <b>Средний возраст:</b> {lesson_plan.study_class.age_group} лет<br/>
        <b>Дата создания:</b> {lesson_plan.created_at.strftime('%d.%m.%Y')}
        """
        story.append(Paragraph(info_text, info_style))
        story.append(Spacer(1, 12))
        
        if lesson_plan.preferences:
            story.append(Paragraph("<b>Пожелания учителя:</b>", styles['Heading2']))
            story.append(Paragraph(lesson_plan.preferences, info_style))
            story.append(Spacer(1, 12))
        
        if lesson_plan.generated_text:
            plain_text = self.markdown_to_plain_text(lesson_plan.generated_text)
            story.append(Paragraph("<b>Содержание плана урока:</b>", styles['Heading2']))
            
            paragraphs = plain_text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    story.append(Paragraph(paragraph.strip(), info_style))
                    story.append(Spacer(1, 6))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def get_pdf_styles(self):
        """Создает стили для PDF с Times New Roman"""
        styles = getSampleStyleSheet()
        
        try:
            styles.add(ParagraphStyle(
                name='TimesNormal',
                fontName='Times-Roman',
                fontSize=12,
                leading=14,
                alignment=TA_JUSTIFY,
            ))
            
            styles.add(ParagraphStyle(
                name='TimesTitle',
                fontName='Times-Bold',
                fontSize=16,
                leading=18,
                alignment=TA_CENTER,
                spaceAfter=30,
            ))
            
            styles.add(ParagraphStyle(
                name='TimesHeading2',
                fontName='Times-Bold',
                fontSize=14,
                leading=16,
                alignment=TA_LEFT,
                spaceBefore=12,
                spaceAfter=6,
            ))
            
            styles['Title'] = styles['TimesTitle']
            styles['Heading2'] = styles['TimesHeading2']
            styles['Normal'] = styles['TimesNormal']
            
        except:
            styles['Title'].alignment = TA_CENTER
            styles['Heading2'].spaceBefore = 12
            styles['Heading2'].spaceAfter = 6
        
        return styles
    
    def create_docx(self, lesson_plan):
        """Создает DOCX файл плана урока"""
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        title = doc.add_heading(f'План урока: {lesson_plan.topic}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Класс: {lesson_plan.study_class.name}')
        doc.add_paragraph(f'Предмет: {lesson_plan.study_class.subject}')
        doc.add_paragraph(f'Количество часов: {lesson_plan.hours}')
        doc.add_paragraph(f'Средний возраст: {lesson_plan.study_class.age_group} лет')
        doc.add_paragraph(f'Дата создания: {lesson_plan.created_at.strftime("%d.%m.%Y")}')
        
        doc.add_paragraph()
        
        if lesson_plan.preferences:
            doc.add_heading('Пожелания учителя:', level=2)
            doc.add_paragraph(lesson_plan.preferences)
            doc.add_paragraph()
        
        if lesson_plan.generated_text:
            doc.add_heading('Содержание плана урока:', level=2)
            plain_text = self.markdown_to_plain_text(lesson_plan.generated_text)
            
            paragraphs = plain_text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style = doc.styles['Normal']
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer